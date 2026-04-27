from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy
import os

doc = Document()

# Page margins
section = doc.sections[0]
section.top_margin = Inches(0.6)
section.bottom_margin = Inches(0.6)
section.left_margin = Inches(0.75)
section.right_margin = Inches(0.75)

# Remove default paragraph spacing
style = doc.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(10)
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.space_after = Pt(0)

def add_horizontal_line(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '2E75B6')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def add_name_header(doc, name):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(name)
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
    return p

def add_contact_line(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    return p

def add_section_heading(doc, title):
    add_horizontal_line(doc)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(title.upper())
    run.bold = True
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
    return p

def add_edu_row(doc, institution, location, degree, dates):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(1)
    r1 = p.add_run(institution)
    r1.bold = True
    r1.font.size = Pt(10)
    tab_run = p.add_run('\t')
    r2 = p.add_run(location)
    r2.font.size = Pt(10)
    p.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), 2)

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(2)
    r3 = p2.add_run(degree)
    r3.italic = True
    r3.font.size = Pt(10)
    tab2 = p2.add_run('\t')
    r4 = p2.add_run(dates)
    r4.font.size = Pt(10)
    p2.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), 2)

def add_job_header(doc, company, dates, title, location):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(1)
    r1 = p.add_run(company)
    r1.bold = True
    r1.font.size = Pt(10)
    p.add_run('\t')
    r2 = p.add_run(dates)
    r2.font.size = Pt(10)
    p.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), 2)

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(2)
    r3 = p2.add_run(title)
    r3.italic = True
    r3.font.size = Pt(10)
    p2.add_run('\t')
    r4 = p2.add_run(location)
    r4.font.size = Pt(10)
    p2.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), 2)

def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.left_indent = Inches(0.2)
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        r1.bold = True
        r1.font.size = Pt(10)
        r2 = p.add_run(text)
        r2.font.size = Pt(10)
    else:
        run = p.add_run(text)
        run.font.size = Pt(10)
    return p

def add_project_header(doc, name, tech, dates):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(1)
    r1 = p.add_run(name + ' ')
    r1.bold = True
    r1.font.size = Pt(10)
    r2 = p.add_run('| ' + tech)
    r2.italic = True
    r2.font.size = Pt(10)
    p.add_run('\t')
    r3 = p.add_run(dates)
    r3.font.size = Pt(10)
    p.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), 2)

def add_skills_row(doc, label, value):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run(label + ': ')
    r1.bold = True
    r1.font.size = Pt(10)
    r2 = p.add_run(value)
    r2.font.size = Pt(10)

# ── HEADER ──────────────────────────────────────────────
add_name_header(doc, 'Muhil Pradhanji')
add_contact_line(doc, '+91 9994811881  |  muhilprathanji225@gmail.com  |  linkedin.com/in/muhil-pradhanji-218524345  |  Chennai, India')

# ── EDUCATION ───────────────────────────────────────────
add_section_heading(doc, 'Education')
add_edu_row(doc,
    'AAA College of Engineering & Technology', 'Sivakasi, India',
    'Bachelor of Computer Science and Engineering (8.01 CGPA)', '2019 – 2023')

# ── EXPERIENCE ──────────────────────────────────────────
add_section_heading(doc, 'Experience')

add_job_header(doc, 'Dhvani Research and Development Solution', 'May 2024 – Present', 'SDE 2', 'Chennai, India')
add_bullet(doc, 'Enhanced the HTTP interceptor to enable customers to intercept, log, and manipulate real-time requests for threat detection.')
add_bullet(doc, 'Refactored the server starter program in Spring Shell to streamline server management and implement health monitoring through JProfiler.')
add_bullet(doc, 'Created a low-code tool using Javassist that automated the generation of Java microservices based on user-provided pseudocode via bytecode manipulation.')

add_job_header(doc, 'Dhvani Research and Development Solution', 'May 2024 – Present', 'Software Developer', 'Chennai, India')
add_bullet(doc, 'Optimized the Mime mail service to handle large attachments by leveraging secondary storage, enabling efficient processing of multi-gigabyte data.')
add_bullet(doc, 'Developed an IP tree and CIDR wild-carding mechanism for efficient and secure IP range-based access control.')
add_bullet(doc, 'Fixed over 50 security vulnerabilities related to regex and datetime handling, significantly improving system security.')

add_job_header(doc, 'AIM Technology', 'June 2023 – Nov. 2023', 'Software Engineer Intern', 'Remote/India')
add_bullet(doc, 'Prepared datasets and trained an ML model to predict inventory requirements for clients like Subway and Cookie Man.')
add_bullet(doc, 'Provided customer support by guiding clients through workarounds, migrations, and issue resolutions.')
add_bullet(doc, 'Designed and implemented a streamlined data pipeline to improve data collection efficiency.')

add_job_header(doc, 'AIM Technology', 'June 2023 – Nov. 2023', 'Software Engineer Intern', 'Remote/India')
add_bullet(doc, 'Developed an intelligent assistive application to help visually challenged individuals identify and locate objects efficiently using smart detection techniques.')
add_bullet(doc, 'Optimized system performance to provide real-time responses, ensuring timely and reliable object identification.')
add_bullet(doc, 'Collaborated on requirement analysis and solution design to align the real-world application.')

# ── PROJECTS ────────────────────────────────────────────
add_section_heading(doc, 'Projects')

add_project_header(doc, 'AI Microservice Generator', 'Java, Spring Boot, Javassist, LLM API (Gemini), Angular', 'March 2025')
add_bullet(doc, 'Built an AI-powered microservice generator that accepts plain-English descriptions and produces production-ready Spring Boot code using Gemini API + prompt engineering.')
add_bullet(doc, 'Extended existing Javassist bytecode manipulation tool to validate and compile LLM-generated Java code at runtime, reducing microservice setup time by ~80%.')
add_bullet(doc, 'Deployed the Angular frontend on Netlify with GitHub CI/CD integration.')

add_project_header(doc, 'RAG Security Vulnerability Chatbot', 'Python, LangChain, ChromaDB, Gemini API, Streamlit', 'April 2025')
add_bullet(doc, 'Built a Retrieval-Augmented Generation (RAG) chatbot ingesting 500+ CVE/OWASP security documents, enabling natural language queries over vulnerability data.')
add_bullet(doc, 'Implemented document chunking, embedding via Gemini text-embedding model, and semantic search using ChromaDB vector database.')
add_bullet(doc, 'Deployed on Streamlit Cloud; leveraged domain knowledge from fixing 50+ real-world security vulnerabilities to curate the knowledge base.')

add_project_header(doc, 'No-Code Platform', 'Go (AST & parser), HTML Widgets', 'January 2024')
add_bullet(doc, 'Created a no-code platform to define referral logic via drag-and-drop pseudo-code integrated via injectable HTML widgets.')

add_project_header(doc, 'Multi-tenant SaaS Boilerplate', 'Angular, Spring Boot, Spring Security', 'April 2024')
add_bullet(doc, 'Developing a boilerplate for rapid SaaS/e-commerce deployment using Temporal to manage distributed workflows.')

add_project_header(doc, 'Kotlin Multiplatform Codebase', 'Kotlin, Android/iOS, MySQL, WASM', 'July 2024')
add_bullet(doc, 'Developed shared logic codebase for mobile platforms with optional WASM support for edge runtime.')

# ── TECHNICAL SKILLS ────────────────────────────────────
add_section_heading(doc, 'Technical Skills')
add_skills_row(doc, 'Languages', 'Java, JavaScript, Python, MySQL, HTML/CSS')
add_skills_row(doc, 'Frameworks', 'Angular, Spring Boot, LangChain, Streamlit')
add_skills_row(doc, 'AI/ML', 'Generative AI, LLM APIs (Gemini, OpenAI), RAG, Prompt Engineering, Vector Databases (ChromaDB)')
add_skills_row(doc, 'Developer Tools', 'VS Code, Visual Studio, GIT, GitHub, Netlify')
add_skills_row(doc, 'Certifications', 'JAVA Certified Developer (2024), MYSQL Certified (2024)')

output_dir = os.path.join(os.getcwd(), 'outputs')
os.makedirs(output_dir, exist_ok=True)
out_path = os.path.join(output_dir, 'Muhil_Pradhanji_Resume.docx')
doc.save(out_path)
print(f"Saved: {out_path}")